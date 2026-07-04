from io import BytesIO

import pytest
from app.file_parser import (
    extract_docx_text,
    extract_resume_text_from_upload,
    normalize_text,
    validate_extracted_text,
)
from docx import Document
from fastapi import HTTPException, UploadFile
from fpdf import FPDF

LONG_TEXT = (
    "Senior Python engineer with FastAPI, SQLAlchemy, and testing experience "
    "delivering resilient backend services across several teams and products."
)


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, text)
    return bytes(pdf.output())


def test_normalize_text_collapses_whitespace() -> None:
    assert normalize_text("  hello   world\n\tfoo ") == "hello world foo"


def test_validate_extracted_text_rejects_short_text() -> None:
    with pytest.raises(HTTPException) as exc_info:
        validate_extracted_text("too short")
    assert exc_info.value.status_code == 422


def test_validate_extracted_text_returns_normalized_when_long_enough() -> None:
    assert validate_extracted_text(f"  {LONG_TEXT}  ") == LONG_TEXT


def test_extract_docx_text_reads_paragraphs_and_tables(tmp_path) -> None:
    document = Document()
    document.add_paragraph(LONG_TEXT)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Cell A"
    table.rows[0].cells[1].text = "Cell B"
    path = tmp_path / "resume.docx"
    document.save(path)

    text = extract_docx_text(path)

    assert LONG_TEXT in text
    assert "Cell A" in text
    assert "Cell B" in text


async def test_extract_resume_text_from_upload_rejects_unsupported_extension() -> None:
    upload = UploadFile(filename="resume.txt", file=BytesIO(b"anything"))
    with pytest.raises(HTTPException) as exc_info:
        await extract_resume_text_from_upload(upload)
    assert exc_info.value.status_code == 415


async def test_extract_resume_text_from_upload_rejects_oversized_file() -> None:
    big = BytesIO(b"a" * (8 * 1024 * 1024 + 1))
    upload = UploadFile(filename="resume.pdf", file=big)
    with pytest.raises(HTTPException) as exc_info:
        await extract_resume_text_from_upload(upload)
    assert exc_info.value.status_code == 413


async def test_extract_resume_text_from_upload_parses_docx() -> None:
    upload = UploadFile(filename="resume.docx", file=BytesIO(_make_docx_bytes([LONG_TEXT])))
    text = await extract_resume_text_from_upload(upload)
    assert LONG_TEXT.split(".")[0] in text


async def test_extract_resume_text_from_upload_parses_pdf() -> None:
    upload = UploadFile(filename="resume.pdf", file=BytesIO(_make_pdf_bytes(LONG_TEXT)))
    text = await extract_resume_text_from_upload(upload)
    assert "Python engineer" in text


async def test_extract_resume_text_from_upload_wraps_parse_errors() -> None:
    upload = UploadFile(filename="resume.pdf", file=BytesIO(b"not a real pdf"))
    with pytest.raises(HTTPException) as exc_info:
        await extract_resume_text_from_upload(upload)
    assert exc_info.value.status_code == 422


async def test_extract_resume_text_from_upload_doc_requires_antiword(monkeypatch) -> None:
    import app.file_parser as file_parser

    monkeypatch.setattr(file_parser.shutil, "which", lambda _: None)
    upload = UploadFile(filename="resume.doc", file=BytesIO(b"legacy doc bytes"))
    with pytest.raises(HTTPException) as exc_info:
        await extract_resume_text_from_upload(upload)
    assert exc_info.value.status_code == 501
