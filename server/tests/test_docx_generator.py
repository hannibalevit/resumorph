from io import BytesIO

import pytest
from app.docx_generator import (
    DOCX_MIME_TYPE,
    render_cover_letter_docx,
    render_resume_docx,
    validate_docx_template,
)
from app.pdf_generator import PDF_MIME_TYPE, render_cover_letter_pdf, render_resume_pdf
from app.schemas import (
    CoverLetter,
    CoverLetterAchievement,
    ResumeCertificationItem,
    ResumeEducationItem,
    ResumeExperienceItem,
    ResumeLanguageItem,
    ResumeProjectItem,
    TailoredResume,
)
from docx import Document
from pydantic import ValidationError
from pypdf import PdfReader


def _resume(**overrides) -> TailoredResume:
    defaults = {
        "candidateName": "Ada Lovelace",
        "contactInfo": "London, UK | ada@example.com | +44 20 1234 5678 | linkedin.com/in/ada",
        "headline": "Senior Backend Engineer",
        "summary": "Experienced Python engineer building resilient services.",
        "competencies": ["Distributed systems", "Python"],
        "skills": ["Python", "FastAPI"],
        "experience": [
            ResumeExperienceItem(
                company="Acme",
                title="Engineer",
                dates="03/2021 - Present",
                bullets=["Built resilient APIs"],
            )
        ],
    }
    defaults.update(overrides)
    return TailoredResume(**defaults)


def _cover_letter(**overrides) -> CoverLetter:
    defaults = {
        "candidateName": "Ada Lovelace",
        "contactInfo": "ada@example.com",
        "roleTitle": "Backend Engineer",
        "company": "Acme",
        "greeting": "Dear Hiring Manager,",
        "opening": "I am excited to apply.",
        "profileIntro": "I bring years of backend experience.",
        "achievements": [CoverLetterAchievement(lead="Scaled APIs", impact="to 1M users.")],
        "closing": "Thank you for your consideration.",
    }
    defaults.update(overrides)
    return CoverLetter(**defaults)


def _paragraphs(data: bytes) -> list[tuple[str, str]]:
    document = Document(BytesIO(data))
    return [(paragraph.style.name, paragraph.text) for paragraph in document.paragraphs]


async def test_render_resume_docx_produces_valid_docx() -> None:
    data, replacements = await render_resume_docx(_resume())

    assert data[:2] == b"PK"
    assert isinstance(replacements, dict)
    assert validate_docx_template(data) == []


async def test_render_resume_pdf_produces_readable_pdf() -> None:
    data, replacements = await render_resume_pdf(_resume())

    assert data.startswith(b"%PDF-")
    assert isinstance(replacements, dict)
    text = "".join(page.extract_text() or "" for page in PdfReader(BytesIO(data)).pages)
    assert "Ada Lovelace" in text
    assert "Work Experience" in text


async def test_render_cover_letter_pdf_produces_readable_pdf() -> None:
    data, _ = await render_cover_letter_pdf(_cover_letter())

    assert data.startswith(b"%PDF-")
    text = "".join(page.extract_text() or "" for page in PdfReader(BytesIO(data)).pages)
    assert "Dear Hiring Manager" in text
    assert "Scaled APIs to 1M users" in text


async def test_render_resume_docx_section_order_and_styles() -> None:
    data, _ = await render_resume_docx(
        _resume(
            projects=[ResumeProjectItem(title="Parser", description="A parser", tech="Python")],
            education=[ResumeEducationItem(institution="MIT", degree="BSc CS", year="2010")],
            certifications=[ResumeCertificationItem(title="AWS SA", org="Amazon", year="2022")],
        )
    )

    styles = [style for style, text in _paragraphs(data) if text.strip()]
    assert styles[0] == "ResumeName"
    assert styles[1] == "ContactLine"
    section_headings = [style for style, text in _paragraphs(data) if style == "SectionHeading"]
    # Standard ATS section order.
    order = [
        "Professional Summary",
        "Core Competencies",
        "Work Experience",
        "Projects",
        "Education",
        "Certifications",
        "Skills",
    ]
    headings_text = [text for style, text in _paragraphs(data) if style == "SectionHeading"]
    assert headings_text == order
    assert section_headings.count("SectionHeading") == len(order)


async def test_render_resume_docx_omits_empty_sections() -> None:
    data, _ = await render_resume_docx(
        _resume(summary="", competencies=[], skills=[], experience=[])
    )

    headings = {text for style, text in _paragraphs(data) if style == "SectionHeading"}
    assert headings == set()
    texts = " ".join(text for _, text in _paragraphs(data))
    assert "Ada Lovelace" in texts


async def test_render_resume_docx_splits_contact_line_by_url() -> None:
    data, _ = await render_resume_docx(
        _resume(contactInfo="Berlin, DE | ada@example.com | linkedin.com/in/ada | github.com/ada")
    )

    contact_lines = [text for style, text in _paragraphs(data) if style == "ContactLine"]
    assert contact_lines[0] == "Berlin, DE | ada@example.com"
    assert contact_lines[1] == "linkedin.com/in/ada | github.com/ada"


async def test_render_resume_docx_sanitizes_unicode_and_reports_counts() -> None:
    data, replacements = await render_resume_docx(
        _resume(summary="Led team — shipped “quality” work fast…")
    )

    texts = " ".join(text for _, text in _paragraphs(data))
    assert "Led team - shipped" in texts
    assert '"quality"' in texts
    assert "..." in texts
    assert replacements["em-dash"] == 1
    assert replacements["smart-double-quote"] == 2
    assert replacements["ellipsis"] == 1


async def test_render_resume_docx_localizes_section_headings() -> None:
    data, _ = await render_resume_docx(_resume(language="de"))

    headings = {text for style, text in _paragraphs(data) if style == "SectionHeading"}
    assert "Berufliches Profil" in headings
    assert "Kernkompetenzen" in headings
    assert "Berufserfahrung" in headings


async def test_render_resume_docx_renders_languages_section() -> None:
    data, _ = await render_resume_docx(
        _resume(
            languages=[
                ResumeLanguageItem(language="English", proficiency="Native"),
                ResumeLanguageItem(language="Spanish", proficiency="Conversational"),
            ]
        )
    )

    headings_text = [text for style, text in _paragraphs(data) if style == "SectionHeading"]
    assert headings_text[-1] == "Languages"
    body_texts = [text for style, text in _paragraphs(data) if style == "BodyText"]
    assert "English (Native), Spanish (Conversational)" in body_texts


async def test_render_resume_docx_omits_languages_section_when_empty() -> None:
    data, _ = await render_resume_docx(_resume(languages=[]))

    headings = {text for style, text in _paragraphs(data) if style == "SectionHeading"}
    assert "Languages" not in headings


async def test_render_resume_docx_letter_page_format() -> None:
    data, _ = await render_resume_docx(_resume(pageFormat="letter"))

    document = Document(BytesIO(data))
    section = document.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert round(section.left_margin.cm, 2) == 1.8


async def test_render_resume_docx_compact_reduces_margins() -> None:
    data, _ = await render_resume_docx(_resume(), compact=True)

    document = Document(BytesIO(data))
    assert round(document.sections[0].left_margin.cm, 2) == 1.5


async def test_render_cover_letter_docx_produces_valid_docx() -> None:
    data, replacements = await render_cover_letter_docx(_cover_letter())

    assert data[:2] == b"PK"
    assert isinstance(replacements, dict)
    assert validate_docx_template(data) == []
    texts = " ".join(text for _, text in _paragraphs(data))
    assert "Ada Lovelace" in texts
    assert "Backend Engineer - Acme" in texts
    assert "Scaled APIs to 1M users." in texts


async def test_render_cover_letter_docx_always_includes_greeting_and_closing() -> None:
    # Regression test: greeting/closing are required schema fields and must never be
    # dropped by a falsy/empty check in the renderer (a prior bug silently skipped a
    # populated closing, making the letter look cut off with no sign-off).
    data, _ = await render_cover_letter_docx(
        _cover_letter(greeting="Hi,", closing="I would welcome the chance to talk further.")
    )

    texts = [text for _, text in _paragraphs(data)]
    assert "Hi," in texts
    assert "I would welcome the chance to talk further." in texts


async def test_render_cover_letter_docx_includes_language_closing_when_present() -> None:
    data, _ = await render_cover_letter_docx(
        _cover_letter(languageClosing="I am also learning Spanish ahead of the move.")
    )

    texts = [text for _, text in _paragraphs(data)]
    assert "I am also learning Spanish ahead of the move." in texts


def test_cover_letter_rejects_empty_closing() -> None:
    with pytest.raises(ValidationError):
        _cover_letter(closing="")


def test_cover_letter_rejects_empty_greeting() -> None:
    with pytest.raises(ValidationError):
        _cover_letter(greeting="")


def test_docx_mime_type() -> None:
    assert DOCX_MIME_TYPE.endswith("wordprocessingml.document")


def test_pdf_mime_type() -> None:
    assert PDF_MIME_TYPE == "application/pdf"
