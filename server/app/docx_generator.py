"""Programmatic .docx rendering for tailored resumes and cover letters.

Builds documents with python-docx only (no hand-written OOXML, no HTML/CSS
intermediate step). The resume shell follows a fixed set of layout rules meant
to keep the output ATS-safe and portable across Word/Google Docs/LibreOffice:
one column, no tables/images/text boxes/floating elements, every paragraph
carries one of a small set of named styles, and section order matches the
standard ATS convention (Professional Summary, Core Competencies, Work
Experience, Projects, Education, Certifications, Skills).
"""

import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import Literal

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Length, Mm, Pt, RGBColor

from app.schemas import CoverLetter, TailoredResume
from app.text_utils import clean_text, headings_for, split_contact_lines

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

PageFormat = Literal["letter", "a4"]


class DocxGenerationError(Exception):
    """Raised when rendering a tailored resume or cover letter to .docx fails."""


# === Page geometry ===========================================================

_PAGE_DIMENSIONS: dict[str, tuple[Length, Length]] = {
    "letter": (Inches(8.5), Inches(11)),
    "a4": (Mm(210), Mm(297)),
}
_MARGIN_NORMAL = Cm(1.8)
_MARGIN_COMPACT = Cm(1.5)


def _configure_section(section, page_format: str, compact: bool) -> Emu:
    """Apply page size/margins and return the usable text width (for tab stops)."""
    width, height = _PAGE_DIMENSIONS.get(page_format, _PAGE_DIMENSIONS["a4"])
    section.page_width = width
    section.page_height = height
    section.orientation = WD_ORIENT.PORTRAIT
    margin = _MARGIN_COMPACT if compact else _MARGIN_NORMAL
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    return Emu(int(width) - 2 * int(margin))


# === Named styles =============================================================

ALLOWED_STYLE_NAMES = {
    "ResumeName",
    "ContactLine",
    "SectionHeading",
    "JobTitle",
    "CompanyLine",
    "BulletItem",
    "BodyText",
    "CompetencyLine",
}


@dataclass(frozen=True)
class _StyleSpec:
    size_pt: float
    bold: bool
    color: str
    space_before_pt: float
    space_after_pt: float
    keep_with_next: bool
    base_style: str | None


_STYLE_SPECS: dict[str, _StyleSpec] = {
    "ResumeName": _StyleSpec(18, True, "000000", 0, 4, False, None),
    "ContactLine": _StyleSpec(10.5, False, "000000", 0, 2, False, None),
    "SectionHeading": _StyleSpec(12, True, "333333", 11, 6, True, None),
    "JobTitle": _StyleSpec(11, True, "000000", 0, 2, True, None),
    "CompanyLine": _StyleSpec(10.5, False, "000000", 0, 4, False, None),
    "BulletItem": _StyleSpec(10.5, False, "000000", 0, 4, False, "List Bullet"),
    "BodyText": _StyleSpec(10.5, False, "000000", 0, 5, False, None),
    "CompetencyLine": _StyleSpec(10.5, False, "000000", 0, 5, False, None),
}


def _spacing(points: float, compact: bool) -> float:
    return round(points * 0.7, 1) if compact else points


def _set_heading_border(style_element) -> None:
    pPr = style_element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "333333")
    border.append(bottom)
    pPr.append(border)


def _get_or_create_paragraph_style(styles, name: str):
    """Return a paragraph style named exactly `name`, reusing a built-in style
    if Word's whitespace-insensitive style-name matching would otherwise
    collide with it (e.g. "BodyText" and the built-in "Body Text" share the
    styleId "BodyText" - adding a second style with that id produces invalid,
    repair-triggering OOXML instead of a distinct style).
    """
    style_id = name.replace(" ", "")
    for style in styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.style_id == style_id:
            style.name = name
            return style
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _configure_styles(document: DocumentType, compact: bool) -> None:
    styles = document.styles
    normal = styles["Normal"]
    # Calibri is the target font; Word/LibreOffice fall back to their own
    # substitute (e.g. Carlito) when Calibri isn't installed, which is the
    # closest OOXML gets to a declared fallback such as Arial.
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.widow_control = True

    for name, spec in _STYLE_SPECS.items():
        style = _get_or_create_paragraph_style(styles, name)
        style.base_style = styles[spec.base_style] if spec.base_style else styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(spec.size_pt)
        style.font.bold = spec.bold
        style.font.color.rgb = RGBColor.from_string(spec.color)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.1
        paragraph_format.space_before = Pt(_spacing(spec.space_before_pt, compact))
        paragraph_format.space_after = Pt(_spacing(spec.space_after_pt, compact))
        paragraph_format.keep_with_next = spec.keep_with_next
        paragraph_format.widow_control = True
        if name == "BulletItem":
            paragraph_format.left_indent = Cm(0.6)
            paragraph_format.first_line_indent = Cm(-0.6)
        if name == "SectionHeading":
            _set_heading_border(style.element)


# === Paragraph builders ========================================================


def _p(document: DocumentType, text: str, style: str, counts: dict[str, int]):
    return document.add_paragraph(clean_text(text, counts), style=style)


def _tab_line(
    document: DocumentType,
    left: str,
    right: str | None,
    style: str,
    usable_width: Emu,
    counts: dict[str, int],
):
    left_clean = clean_text(left, counts)
    right_clean = clean_text(right, counts) if right else ""
    paragraph = document.add_paragraph(style=style)
    if right_clean:
        paragraph.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
        paragraph.add_run(f"{left_clean}\t{right_clean}")
    else:
        paragraph.add_run(left_clean)
    return paragraph


# === Resume ====================================================================


def _build_resume_document(
    resume: TailoredResume, *, compact: bool = False
) -> tuple[bytes, dict[str, int]]:
    counts: dict[str, int] = {}
    document = Document()
    _configure_styles(document, compact)
    usable_width = _configure_section(document.sections[0], resume.page_format, compact)
    headings = headings_for(resume.language)

    _p(document, resume.candidate_name, "ResumeName", counts)

    primary_line, url_line = split_contact_lines(resume.contact_info or "")
    if primary_line:
        _p(document, primary_line, "ContactLine", counts)
    if url_line:
        _p(document, url_line, "ContactLine", counts)

    if resume.summary.strip():
        _p(document, headings["summary"], "SectionHeading", counts)
        _p(document, resume.summary, "BodyText", counts)

    competencies = [item.strip() for item in resume.competencies if item.strip()]
    if competencies:
        _p(document, headings["competencies"], "SectionHeading", counts)
        _p(document, "  |  ".join(competencies), "CompetencyLine", counts)

    if resume.experience:
        _p(document, headings["experience"], "SectionHeading", counts)
        for job in resume.experience:
            _p(document, job.title, "JobTitle", counts)
            _tab_line(document, job.company, job.dates, "CompanyLine", usable_width, counts)
            if job.location:
                _p(document, job.location, "BodyText", counts)
            for bullet in job.bullets:
                if bullet.strip():
                    _p(document, bullet.strip(), "BulletItem", counts)

    if resume.projects:
        _p(document, headings["projects"], "SectionHeading", counts)
        for project in resume.projects:
            title = f"{project.title} ({project.badge})" if project.badge else project.title
            _p(document, title, "JobTitle", counts)
            if project.description:
                _p(document, project.description, "BodyText", counts)
            if project.tech:
                _p(document, project.tech, "BodyText", counts)

    if resume.education:
        _p(document, headings["education"], "SectionHeading", counts)
        for edu in resume.education:
            _tab_line(document, edu.degree, edu.year, "JobTitle", usable_width, counts)
            _p(document, edu.institution, "CompanyLine", counts)
            if edu.description:
                _p(document, edu.description, "BodyText", counts)

    if resume.certifications:
        _p(document, headings["certifications"], "SectionHeading", counts)
        for cert in resume.certifications:
            label = f"{cert.title} - {cert.org}" if cert.org else cert.title
            _tab_line(document, label, cert.year, "CompanyLine", usable_width, counts)

    skills = [item.strip() for item in resume.skills if item.strip()]
    if skills:
        _p(document, headings["skills"], "SectionHeading", counts)
        _p(document, ", ".join(skills), "BodyText", counts)

    languages = [item for item in resume.languages if item.language.strip()]
    if languages:
        _p(document, headings["languages"], "SectionHeading", counts)
        formatted_languages = [
            f"{item.language.strip()} ({item.proficiency.strip()})"
            if item.proficiency and item.proficiency.strip()
            else item.language.strip()
            for item in languages
        ]
        _p(document, ", ".join(formatted_languages), "BodyText", counts)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), counts


async def render_resume_docx(
    resume: TailoredResume, *, compact: bool = False
) -> tuple[bytes, dict[str, int]]:
    try:
        return _build_resume_document(resume, compact=compact)
    except Exception as exc:
        raise DocxGenerationError(f"Failed to render resume DOCX: {exc}") from exc


# === Cover letter ==============================================================


def _build_cover_letter_document(
    letter: CoverLetter, *, compact: bool = False
) -> tuple[bytes, dict[str, int]]:
    counts: dict[str, int] = {}
    document = Document()
    _configure_styles(document, compact)
    _configure_section(document.sections[0], letter.page_format, compact)

    _p(document, letter.candidate_name, "ResumeName", counts)

    primary_line, url_line = split_contact_lines(letter.contact_info or "")
    if primary_line:
        _p(document, primary_line, "ContactLine", counts)
    if url_line:
        _p(document, url_line, "ContactLine", counts)

    credentials = [item.strip() for item in letter.credentials if item.strip()]
    if credentials:
        _p(document, " | ".join(credentials), "ContactLine", counts)

    role_line = letter.role_title
    if letter.company:
        role_line = f"{role_line} - {letter.company}"
    _p(document, role_line, "JobTitle", counts)

    if letter.dateline and letter.dateline.strip():
        _p(document, letter.dateline, "BodyText", counts)
    # greeting/opening/profileIntro/closing are required schema fields (min_length-enforced),
    # so they always render - unlike the truly optional fields below, they must never be
    # skipped on a falsy/empty check (that previously dropped a populated closing silently).
    _p(document, letter.greeting, "BodyText", counts)
    _p(document, letter.opening, "BodyText", counts)
    _p(document, letter.profile_intro, "BodyText", counts)

    for item in letter.achievements:
        _p(document, f"{item.lead} {item.impact}", "BulletItem", counts)

    if letter.problems and letter.problems.strip():
        _p(document, letter.problems, "BodyText", counts)
    _p(document, letter.closing, "BodyText", counts)
    if letter.language_closing and letter.language_closing.strip():
        _p(document, letter.language_closing, "BodyText", counts)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), counts


async def render_cover_letter_docx(
    letter: CoverLetter, *, compact: bool = False
) -> tuple[bytes, dict[str, int]]:
    try:
        return _build_cover_letter_document(letter, compact=compact)
    except Exception as exc:
        raise DocxGenerationError(f"Failed to render cover letter DOCX: {exc}") from exc


# === Structural validation =====================================================
# Adaptation of the spec's "print a checklist" step for a library used inside a
# live request path: returns violations instead of printing, so callers/tests
# can assert the template stays compliant as it evolves.


def validate_docx_template(data: bytes) -> list[str]:
    """Check the generated .docx against the template's hard layout rules.

    Returns an empty list when the document is fully compliant, otherwise one
    message per violated rule (zero tables/images/text boxes, no header/footer
    content, a single single-column section, and every paragraph using one of
    the template's named styles).
    """
    violations: list[str] = []
    with zipfile.ZipFile(BytesIO(data)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")

    if "<w:tbl>" in xml or "<w:tbl " in xml or "<w:tbl/>" in xml:
        violations.append("document contains a table")
    if "<w:drawing" in xml or "<w:pict" in xml or "<w:object" in xml:
        violations.append("document contains an image, shape, or embedded object")
    if "<w:txbxContent" in xml:
        violations.append("document contains a text box")
    if re.search(r'<w:cols\b[^>]*\bw:num="(?!1")\d', xml):
        violations.append("document uses more than one column")

    document = Document(BytesIO(data))
    if len(document.sections) != 1:
        violations.append("document has more than one section")

    for section in document.sections:
        header_text = "".join(paragraph.text for paragraph in section.header.paragraphs).strip()
        footer_text = "".join(paragraph.text for paragraph in section.footer.paragraphs).strip()
        if header_text or footer_text:
            violations.append("document has header/footer content")
            break

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else None
        if paragraph.text.strip() and style_name not in ALLOWED_STYLE_NAMES:
            violations.append(f"paragraph uses non-template style '{style_name}'")
            break

    return violations
