from io import BytesIO

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from app.schemas import LegacyTailoredResume

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)


def add_bullets(document: Document, values: list[str]) -> None:
    for value in values:
        if value.strip():
            document.add_paragraph(value.strip(), style="List Bullet")


def create_docx_resume(resume: LegacyTailoredResume) -> bytes:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    name = document.add_paragraph()
    name_run = name.add_run(resume.candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(16)

    if resume.contact_info:
        contact = document.add_paragraph()
        contact.add_run(resume.contact_info)

    headline = document.add_paragraph()
    headline.add_run(resume.headline)

    add_heading(document, "Summary")
    document.add_paragraph(resume.summary)

    if resume.skills:
        add_heading(document, "Skills")
        document.add_paragraph(", ".join(skill.strip() for skill in resume.skills if skill.strip()))

    if resume.experience:
        add_heading(document, "Experience")
        for index, item in enumerate(resume.experience):
            if index > 0:
                spacer = document.add_paragraph()
                spacer.add_run().add_break(WD_BREAK.LINE)

            title_line = f"{item.title} - {item.company}"
            if item.dates:
                title_line = f"{title_line} | {item.dates}"

            role = document.add_paragraph()
            role_run = role.add_run(title_line)
            role_run.bold = True
            add_bullets(document, item.bullets)

    if resume.education:
        add_heading(document, "Education")
        add_bullets(document, resume.education)

    if resume.projects:
        add_heading(document, "Projects")
        add_bullets(document, resume.projects)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_docx_text(title: str, body: str) -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10.5)
    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    for paragraph in body.split("\n\n"):
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
