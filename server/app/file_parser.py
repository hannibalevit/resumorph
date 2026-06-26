import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

MAX_UPLOAD_BYTES = 8 * 1024 * 1024
MIN_EXTRACTED_TEXT_LENGTH = 100
SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}


def normalize_text(value: str) -> str:
    return " ".join(value.split()).strip()


def validate_extracted_text(text: str) -> str:
    normalized_text = normalize_text(text)

    if len(normalized_text) < MIN_EXTRACTED_TEXT_LENGTH:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract enough text from this resume file.",
        )

    return normalized_text


def extract_pdf_text(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    page_texts = [page.extract_text() or "" for page in reader.pages]
    return validate_extracted_text("\n".join(page_texts))


def extract_docx_text(file_path: Path) -> str:
    document = Document(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)

    return validate_extracted_text("\n".join(paragraphs))


def extract_doc_text(file_path: Path) -> str:
    antiword = shutil.which("antiword")
    if not antiword:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="Legacy .doc parsing requires antiword on the backend.",
        )

    completed = subprocess.run(
        [antiword, str(file_path)],
        check=False,
        capture_output=True,
        text=True,
        timeout=20,
    )

    if completed.returncode != 0:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract text from the .doc resume file.",
        )

    return validate_extracted_text(completed.stdout)


async def extract_resume_text_from_upload(file: UploadFile) -> str:
    filename = Path(file.filename or "").name
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Supported backend parsing formats are .pdf, .doc, and .docx.",
        )

    contents = await file.read()
    if len(contents) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="Resume file is too large. Maximum supported upload size is 8 MB.",
        )

    with tempfile.TemporaryDirectory() as temporary_directory:
        file_path = Path(temporary_directory) / f"resume{extension}"
        file_path.write_bytes(contents)

        try:
            if extension == ".pdf":
                return extract_pdf_text(file_path)
            if extension == ".docx":
                return extract_docx_text(file_path)
            return extract_doc_text(file_path)
        except HTTPException:
            raise
        except subprocess.TimeoutExpired as exc:
            raise HTTPException(
                status_code=status.HTTP_504_GATEWAY_TIMEOUT,
                detail="Timed out while parsing the resume file.",
            ) from exc
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not parse the resume file.",
            ) from exc
